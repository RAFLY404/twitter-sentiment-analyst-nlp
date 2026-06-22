from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "IEEE_PPKM_Sentiment_Analysis_Report.docx"
FIGURE_DIR = ROOT / ".codex" / "verification"

BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "E7E6E6"
WHITE = "FFFFFF"

# The documents skill requires a resolved preset. This report uses the
# compact_reference_guide preset as its density baseline, with one named
# override: IEEE Conference Paper. The override supplies A4 geometry,
# Times New Roman typography, a single-column title block, and a two-column body.
TOKENS = {
    "page_width": Cm(21.0),
    "page_height": Cm(29.7),
    "margin_top": Cm(1.65),
    "margin_bottom": Cm(1.65),
    "margin_left": Cm(1.60),
    "margin_right": Cm(1.60),
    "column_gap_dxa": "360",
    "body_font": "Times New Roman",
    "body_size": 10,
    "body_after": 3,
    "table_width_dxa": 4800,
    "table_indent_dxa": 0,
    "cell_margin_dxa": 70,
}


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="808080", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent_dxa=0):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, *(TOKENS["cell_margin_dxa"],) * 4)


def set_columns(section, count=2, space_dxa="360"):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    for child in list(cols):
        cols.remove(child)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_dxa))
    cols.set(qn("w:equalWidth"), "1")


def configure_section(section, columns=1):
    section.page_width = TOKENS["page_width"]
    section.page_height = TOKENS["page_height"]
    section.top_margin = TOKENS["margin_top"]
    section.bottom_margin = TOKENS["margin_bottom"]
    section.left_margin = TOKENS["margin_left"]
    section.right_margin = TOKENS["margin_right"]
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)
    set_columns(section, columns, TOKENS["column_gap_dxa"])


def set_font(run, size=None, bold=None, italic=None, color=BLACK, small_caps=False):
    run.font.name = TOKENS["body_font"]
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), TOKENS["body_font"])
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), TOKENS["body_font"])
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), TOKENS["body_font"])
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color
    run.font.small_caps = small_caps


def style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3, line=1.0, first_line=True):
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = True
    pf.widow_control = True
    if first_line:
        pf.first_line_indent = Inches(0.14)
    else:
        pf.first_line_indent = Inches(0)


def add_body(doc, text, *, after=3, first_line=True):
    p = doc.add_paragraph()
    style_paragraph(p, after=after, first_line=first_line)
    run = p.add_run(text)
    set_font(run, size=TOKENS["body_size"])
    return p


def add_ieee_heading(doc, number, title):
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=7, after=3, first_line=False)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{number}. {title.upper()}")
    set_font(run, size=10, small_caps=True)
    return p


def add_subheading(doc, letter, title):
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=2, first_line=False)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{letter}. {title}")
    set_font(run, size=10, italic=True)
    return p


def add_equation(doc, equation, number):
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3, first_line=False)
    r = p.add_run(equation)
    set_font(r, size=9, italic=True)
    r2 = p.add_run(f"    ({number})")
    set_font(r2, size=9)
    return p


def add_table_caption(doc, number, title):
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, first_line=False)
    p.paragraph_format.keep_with_next = True
    r1 = p.add_run(f"TABLE {number}\n")
    set_font(r1, size=8, small_caps=True)
    r2 = p.add_run(title.upper())
    set_font(r2, size=8, small_caps=True)
    return p


def add_data_table(doc, headers, rows, widths, font_size=7.5, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, TOKENS["table_indent_dxa"])
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0, first_line=False)
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True)

    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, alignment=alignment, after=0, first_line=False)
            r = p.add_run(str(value))
            set_font(r, size=font_size)

    # Geometry must be re-applied after all rows have been added.
    set_table_geometry(table, widths, TOKENS["table_indent_dxa"])
    return table


def add_figure(doc, path, caption, width_inches=3.15):
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=1, first_line=False)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))

    cap = doc.add_paragraph()
    style_paragraph(cap, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4, first_line=False)
    r = cap.add_run(caption)
    set_font(r, size=8)
    return cap


def add_reference(doc, number, text):
    p = doc.add_paragraph()
    style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=1.5, line=1.0, first_line=False)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    r = p.add_run(f"[{number}] {text}")
    set_font(r, size=8)
    return p


def build_document():
    doc = Document()
    doc.core_properties.title = "Comparative Evaluation of Classical Machine Learning Models for Indonesian PPKM Tweet Sentiment Analysis Using TF-IDF"
    doc.core_properties.subject = "IEEE-style NLP course project report"
    doc.core_properties.author = "NLP Project Group"
    doc.core_properties.keywords = "sentiment analysis, Indonesian NLP, PPKM, TF-IDF, logistic regression"

    configure_section(doc.sections[0], columns=1)

    normal = doc.styles["Normal"]
    normal.font.name = TOKENS["body_font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["body_font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["body_font"])
    normal.font.size = Pt(TOKENS["body_size"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(TOKENS["body_after"])
    normal.paragraph_format.line_spacing = 1.0

    title = doc.add_paragraph()
    style_paragraph(title, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=7, first_line=False)
    title.paragraph_format.keep_with_next = True
    tr = title.add_run("Comparative Evaluation of Classical Machine Learning Models for Indonesian PPKM Tweet Sentiment Analysis Using TF-IDF")
    set_font(tr, size=19, bold=False)

    affiliation = "Computer Science Department\nSchool of Computer Science\nBina Nusantara University\nJakarta, Indonesia"
    author_rows = [
        ["Nama Anggota 1", "Nama Anggota 2", "Nama Anggota 3"],
        ["Nama Anggota 4", "Nama Anggota 5", ""],
    ]
    author_table = doc.add_table(rows=2, cols=3)
    author_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    author_table.autofit = False
    set_table_geometry(author_table, [3360, 3360, 3360], 0)
    # Borderless author block, as in an IEEE conference template.
    tbl_pr = author_table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)

    for row_i, row in enumerate(author_rows):
        for col_i, name in enumerate(row):
            cell = author_table.cell(row_i, col_i)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0, first_line=False)
            if not name:
                continue
            nr = p.add_run(name + "\n")
            set_font(nr, size=9, color=MUTED)
            ar = p.add_run(affiliation + "\n")
            set_font(ar, size=8)
            er = p.add_run(f"nama{row_i * 3 + col_i + 1}@binus.ac.id")
            set_font(er, size=8, color=MUTED)

    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=2, first_line=False)

    abstract = doc.add_paragraph()
    style_paragraph(abstract, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=3, first_line=False)
    label = abstract.add_run("Abstract—")
    set_font(label, size=9, bold=True, italic=True)
    text = abstract.add_run(
        "Social-media discussions around public policy contain large volumes of noisy and informal Indonesian text. "
        "This study evaluates four classical supervised classifiers for three-class sentiment analysis of tweets related to "
        "Pemberlakuan Pembatasan Kegiatan Masyarakat (PPKM). The pipeline normalizes URLs, mentions, hashtags, repeated "
        "characters, symbols, and Indonesian slang before representing tweets with TF-IDF unigram and bigram features. "
        "Logistic Regression, Linear Support Vector Classification, Multinomial Naive Bayes, and Complement Naive Bayes "
        "are compared on the same stratified 80/20 split of 23,644 labeled tweets. Although Linear SVC achieves the highest "
        "accuracy (0.8579), Logistic Regression provides the best Macro F1 (0.7373) and is therefore selected under the "
        "project's class-balanced criterion. The selected model obtains 0.8414 accuracy and 0.8499 Weighted F1. Per-class "
        "analysis shows strong performance for neutral and negative tweets but lower F1 for the minority positive class. "
        "The results demonstrate that a reproducible TF-IDF pipeline remains an effective and interpretable baseline for "
        "Indonesian policy-related sentiment analysis, while also exposing the limitations caused by class imbalance, "
        "domain specificity, and context-insensitive sparse features."
    )
    set_font(text, size=9, italic=True)

    keywords = doc.add_paragraph()
    style_paragraph(keywords, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=5, first_line=False)
    kr = keywords.add_run("Keywords—")
    set_font(kr, size=9, bold=True, italic=True)
    kv = keywords.add_run("Indonesian NLP, sentiment analysis, PPKM, TF-IDF, Logistic Regression, Linear SVC, class imbalance")
    set_font(kv, size=9, italic=True)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, columns=2)

    add_ieee_heading(doc, "I", "Introduction")
    add_body(
        doc,
        "Social-media platforms provide a direct but noisy channel for observing public reactions to government policy. "
        "During PPKM, Indonesian tweets expressed support, criticism, personal experience, and factual updates in highly "
        "informal language. Manually reading this volume of text is impractical; sentiment analysis converts the unstructured "
        "stream into positive, neutral, and negative categories that can support aggregate analysis [1]. Earlier work has "
        "shown that sparse lexical features combined with supervised classifiers form competitive and interpretable baselines "
        "for sentiment classification [2]."
    )
    add_body(
        doc,
        "Indonesian Twitter data adds several challenges. Tweets contain abbreviations, slang, user mentions, links, hashtags, "
        "creative spelling, and repeated characters. The available PPKM corpus is also substantially imbalanced: neutral tweets "
        "represent nearly three quarters of the observations, while positive tweets form the smallest class. Consequently, a "
        "high accuracy score can hide weak minority-class performance. This project therefore treats Macro F1 as the primary "
        "selection metric and uses accuracy only as a tiebreaker."
    )
    add_body(
        doc,
        "This paper contributes a reproducible comparison of four classical machine-learning pipelines under identical data, "
        "preprocessing, feature, and split conditions. It also reports class-level errors and links the experimental results to "
        "the deployable command-line prediction artifact included in the project repository. The scope is deliberately limited "
        "to PPKM tweet sentiment; extractive summarization demonstrations in the notebook are not evaluated in this study."
    )

    add_ieee_heading(doc, "II", "Related Work")
    add_body(
        doc,
        "Sentiment analysis combines natural-language processing and statistical learning to infer subjective orientation from "
        "text [1]. The classic comparison by Pang et al. [2] established machine-learning methods as strong baselines for "
        "document-level sentiment classification. For Indonesian social media, normalization is especially important because "
        "lexical variation fragments the feature space. The slang resource used by this project originates from the Indonesian "
        "abusive-language dataset released by Ibrohim and Budi [4]."
    )
    add_body(
        doc,
        "TF-IDF remains widely used for text categorization because it emphasizes terms that are frequent in a document but "
        "not ubiquitous across the collection [5], [6]. Linear classifiers are well matched to the resulting high-dimensional "
        "sparse vectors [7], [8], while Naive Bayes provides a computationally inexpensive probabilistic baseline [9]. The "
        "implementation uses scikit-learn [10], allowing all four classifiers to share the same vectorizer and evaluation code."
    )

    add_ieee_heading(doc, "III", "Methodology")
    add_subheading(doc, "A", "Dataset and Label Distribution")
    add_body(
        doc,
        "The experiment uses the Indonesian Twitter Sentiment Analysis Dataset - PPKM [3]. After removing rows with missing "
        "text or labels and discarding empty normalized outputs, 23,644 tweets remain. The original integer mapping is 0 for "
        "positive, 1 for neutral, and 2 for negative. Table I shows the strong class imbalance."
    )
    add_table_caption(doc, "I", "Distribution of PPKM Sentiment Labels")
    add_data_table(
        doc,
        ["Class", "Count", "Share"],
        [
            ["Positive", "1,958", "8.28%"],
            ["Neutral", "17,706", "74.88%"],
            ["Negative", "3,980", "16.83%"],
        ],
        [2100, 1350, 1350],
        numeric_cols=(1, 2),
    )

    add_subheading(doc, "B", "Text Preprocessing")
    add_body(
        doc,
        "The normalization function first converts text to lowercase and decodes HTML entities. URLs are deleted, user mentions "
        "are replaced by the token 'user', and the hash symbol is removed while retaining the hashtag word. Sequences of three "
        "or more identical characters are reduced to two characters. Non-alphanumeric symbols are removed, whitespace is "
        "collapsed, and every remaining token is replaced with its formal equivalent when it appears in the slang dictionary. "
        "This sequence reduces feature fragmentation while preserving topical words and the presence of mentions."
    )

    add_subheading(doc, "C", "TF-IDF Representation")
    add_body(
        doc,
        "Each normalized tweet is converted to a sparse TF-IDF vector. The vectorizer includes unigrams and bigrams, removes "
        "a compact Indonesian stop-word list, excludes terms observed in fewer than two documents or more than 95% of the "
        "corpus, and caps the vocabulary at 30,000 features. Sublinear term frequency reduces the influence of repeated tokens. "
        "With scikit-learn's default smoothed inverse document frequency, the unnormalized weight is:"
    )
    add_equation(doc, "w(t,d) = [1 + log tf(t,d)] [log((1 + N)/(1 + df(t))) + 1]", 1)
    add_body(
        doc,
        "where tf(t,d) is the frequency of term t in tweet d, df(t) is its document frequency, and N is the number of training "
        "tweets. The final vector is L2-normalized. Bigrams such as 'tidak setuju' or 'masyarakat kecil' retain local context that "
        "would be lost if only independent words were used."
    )

    add_subheading(doc, "D", "Compared Classifiers")
    add_body(
        doc,
        "Four classifiers are evaluated. Logistic Regression uses the lbfgs solver, a maximum of 1,000 iterations, and balanced "
        "class weights. Linear SVC also uses balanced class weights and a maximum of 5,000 iterations. Multinomial Naive Bayes "
        "and Complement Naive Bayes both use alpha = 0.5. Balanced class weights increase the contribution of underrepresented "
        "classes during optimization, whereas Complement Naive Bayes is designed to be more robust than standard Multinomial "
        "Naive Bayes on imbalanced text data."
    )

    add_subheading(doc, "E", "Evaluation Protocol")
    add_body(
        doc,
        "A single stratified split allocates 18,915 tweets to training and 4,729 to testing. The test proportion is 20%, the "
        "random seed is 42, and the same split is reused for every model. Stratification preserves the original label proportions. "
        "Accuracy, Macro F1, Weighted F1, a confusion matrix, and a per-class classification report are computed. Macro F1 "
        "assigns equal weight to each class and is therefore more informative than accuracy alone under imbalance [11], [12]."
    )

    add_ieee_heading(doc, "IV", "Results and Discussion")
    add_subheading(doc, "A", "Overall Model Comparison")
    add_table_caption(doc, "II", "Test-Set Performance of the Compared Models")
    add_data_table(
        doc,
        ["Model", "Acc.", "Macro F1", "Wtd. F1"],
        [
            ["Logistic Regression", "0.8414", "0.7373", "0.8499"],
            ["Linear SVC", "0.8579", "0.7337", "0.8568"],
            ["Multinomial NB", "0.8450", "0.6597", "0.8319"],
            ["Complement NB", "0.7919", "0.6716", "0.8048"],
        ],
        [1900, 950, 1000, 950],
        font_size=7.1,
        numeric_cols=(1, 2, 3),
    )
    add_body(
        doc,
        "Linear SVC produces the highest accuracy and Weighted F1, but its Macro F1 is 0.0036 below Logistic Regression. The "
        "difference indicates that Linear SVC benefits slightly more from the dominant neutral class, while Logistic Regression "
        "achieves the best average balance across positive, neutral, and negative labels. Because the selection rule prioritizes "
        "Macro F1, Logistic Regression is saved as the deployment pipeline. Both Naive Bayes variants trail the linear "
        "discriminative models, especially on Macro F1."
    )
    add_figure(doc, FIGURE_DIR / "model_comparison.png", "Fig. 1. Accuracy, Macro F1, and Weighted F1 for the four evaluated classifiers.", 3.15)

    add_subheading(doc, "B", "Per-Class Performance")
    add_table_caption(doc, "III", "Classification Report for the Selected Logistic Regression Model")
    add_data_table(
        doc,
        ["Class", "Prec.", "Recall", "F1", "Support"],
        [
            ["Positive", "0.49", "0.67", "0.56", "392"],
            ["Neutral", "0.95", "0.86", "0.91", "3,541"],
            ["Negative", "0.67", "0.83", "0.74", "796"],
        ],
        [1200, 850, 850, 800, 1100],
        font_size=7.2,
        numeric_cols=(1, 2, 3, 4),
    )
    add_body(
        doc,
        "Neutral tweets are classified most reliably, with precision 0.95 and F1 0.91. Negative tweets obtain recall 0.83 and "
        "F1 0.74, indicating that negative lexical cues are detected reasonably well. Positive sentiment remains the main "
        "weakness: recall reaches 0.67, but precision is only 0.49, so many tweets predicted as positive belong to another class. "
        "The minority size of the positive class and overlap between positive statements and neutral information are plausible "
        "causes."
    )
    add_figure(doc, FIGURE_DIR / "confusion_matrix.png", "Fig. 2. Confusion matrix for the selected Logistic Regression pipeline (rows: true labels; columns: predicted labels).", 3.05)

    add_subheading(doc, "C", "Error Patterns and Metric Interpretation")
    add_body(
        doc,
        "The confusion matrix shows 3,059 correct neutral predictions, 657 correct negative predictions, and 263 correct "
        "positive predictions. The largest off-diagonal errors are neutral-to-negative (260) and neutral-to-positive (222), "
        "which reflects both the size of the neutral class and lexical overlap in policy discussions. The model also maps 83 "
        "negative tweets to neutral. These errors are consistent with a bag-of-ngrams representation: negation phrases may be "
        "captured, but sarcasm, long-distance context, stance, and implicit sentiment are not modeled directly."
    )
    add_body(
        doc,
        "The comparison also demonstrates why a single metric is insufficient. Reporting only 0.8579 accuracy would favor "
        "Linear SVC, whereas the project's explicit emphasis on equal class treatment favors Logistic Regression. Weighted F1 "
        "tracks accuracy closely because it is dominated by neutral support. Macro F1 therefore gives a clearer view of whether "
        "the system remains useful for the smaller positive and negative categories."
    )

    add_ieee_heading(doc, "V", "Implementation and Reproducibility")
    add_body(
        doc,
        "The complete workflow is implemented in src/train_sentiment.py. A scikit-learn Pipeline binds preprocessing-compatible "
        "TF-IDF features to each classifier, preventing accidental differences in feature preparation. The script prints the "
        "model comparison, exports plots and a CSV table, and serializes the selected pipeline together with label metadata and "
        "the slang map to models/ppkm_sentiment_pipeline.pkl. The companion src/predict_sentiment.py script loads the artifact "
        "and supports either a single text input or batch prediction from a CSV file."
    )
    add_body(
        doc,
        "Reproducibility is supported by a fixed random state, a documented dependency list, a deterministic stratified split, "
        "and the same vectorization settings for all classifiers. Running 'python src/train_sentiment.py' regenerates the report "
        "figures and model comparison. Running 'python src/predict_sentiment.py <text>' returns the predicted class; probability "
        "scores are reported when the selected classifier implements predict_proba."
    )

    add_ieee_heading(doc, "VI", "Limitations and Future Work")
    add_body(
        doc,
        "First, the corpus is tied to PPKM discourse, so performance cannot be assumed to transfer to products, elections, or "
        "general Indonesian conversation. Second, the study uses one train/test split rather than repeated cross-validation. "
        "Third, no systematic hyperparameter search, probability calibration, stemming, or learned embedding is evaluated. "
        "Fourth, the slang dictionary is static and may not cover new forms, code-switching, or context-dependent meanings. "
        "Finally, TF-IDF cannot represent sarcasm, implicit stance, or compositional semantics beyond short n-grams."
    )
    add_body(
        doc,
        "Future work should use stratified cross-validation and report confidence intervals; tune regularization, n-gram range, "
        "and class weighting; inspect misclassified samples by class; and compare the classical baseline with Indonesian "
        "transformer encoders. Data-level balancing or targeted augmentation for positive tweets should be evaluated carefully "
        "to avoid distribution shift. External evaluation on a different Indonesian policy dataset would provide the strongest "
        "test of generalization."
    )

    add_ieee_heading(doc, "VII", "Conclusion")
    add_body(
        doc,
        "This study compared four classical classifiers for sentiment analysis of 23,644 Indonesian PPKM tweets under a shared "
        "TF-IDF pipeline. Linear SVC achieved the highest accuracy, but Logistic Regression achieved the best Macro F1 and was "
        "selected according to the project's class-balanced criterion. Its 0.8414 accuracy, 0.7373 Macro F1, and 0.8499 Weighted "
        "F1 establish a strong, interpretable baseline. The per-class results also make the remaining problem clear: the neutral "
        "class is modeled well, while positive sentiment requires better representation and more balanced evidence. The project "
        "therefore succeeds as a reproducible classical NLP benchmark and provides a defensible foundation for future comparison "
        "with contextual Indonesian language models."
    )

    add_ieee_heading(doc, "", "References")
    # Remove the leading dot added by the generic heading function.
    ref_heading = doc.paragraphs[-1]
    ref_heading.runs[0].text = "REFERENCES"

    references = [
        "B. Liu, Sentiment Analysis and Opinion Mining. San Rafael, CA, USA: Morgan & Claypool, 2012.",
        "B. Pang, L. Lee, and S. Vaithyanathan, \"Thumbs up? Sentiment classification using machine learning techniques,\" in Proc. EMNLP, 2002, pp. 79-86.",
        "A. Purnama, \"Indonesian Twitter Sentiment Analysis Dataset - PPKM,\" Kaggle. [Online]. Available: https://www.kaggle.com/datasets/anggapurnama/twitter-dataset-ppkm. Accessed: Jun. 22, 2026.",
        "M. O. Ibrohim and I. Budi, \"Multi-label hate speech and abusive language detection in Indonesian Twitter,\" in Proc. Third Workshop on Abusive Language Online, 2019, pp. 46-57.",
        "C. D. Manning, P. Raghavan, and H. Schutze, Introduction to Information Retrieval. Cambridge, U.K.: Cambridge Univ. Press, 2008.",
        "J. Ramos, \"Using TF-IDF to determine word relevance in document queries,\" in Proc. 1st Instructional Conf. Machine Learning, 2003, pp. 133-142.",
        "Y. Yang and X. Liu, \"A re-examination of text categorization methods,\" in Proc. 22nd Annu. Int. ACM SIGIR Conf., 1999, pp. 42-49.",
        "R.-E. Fan, K.-W. Chang, C.-J. Hsieh, X.-R. Wang, and C.-J. Lin, \"LIBLINEAR: A library for large linear classification,\" J. Mach. Learn. Res., vol. 9, pp. 1871-1874, 2008.",
        "A. McCallum and K. Nigam, \"A comparison of event models for Naive Bayes text classification,\" in AAAI-98 Workshop on Learning for Text Categorization, 1998, pp. 41-48.",
        "F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" J. Mach. Learn. Res., vol. 12, pp. 2825-2830, 2011.",
        "T. Saito and M. Rehmsmeier, \"The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets,\" PLoS ONE, vol. 10, no. 3, e0118432, 2015.",
        "N. Japkowicz and S. Stephen, \"The class imbalance problem: A systematic study,\" Intell. Data Anal., vol. 6, no. 5, pp. 429-449, 2002.",
    ]
    for index, reference in enumerate(references, start=1):
        add_reference(doc, index, reference)

    # Explicitly keep headers and footers empty, matching IEEE manuscript style.
    for section in doc.sections:
        for p in section.header.paragraphs + section.footer.paragraphs:
            p.clear()

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
